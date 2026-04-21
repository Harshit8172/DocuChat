import os
import io
import csv
import traceback
from datetime import datetime, timedelta
import secrets

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, Header, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import List, Optional

import pypdf
import docx
import pandas as pd
import jwt

# --- BCRYPT PASSLIB BUG FIX ---
# Modern bcrypt removed __about__ which breaks passlib. We patch it here.
import bcrypt
class _BcryptAbout:
    __version__ = "3.2.0"
bcrypt.__about__ = _BcryptAbout

from passlib.context import CryptContext

from sqlalchemy import create_engine, Column, Integer, String, ForeignKey
from sqlalchemy.orm import declarative_base, sessionmaker, relationship, Session

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEndpointEmbeddings
from langchain_openai import ChatOpenAI
from langchain_pinecone import PineconeVectorStore
from pinecone import Pinecone

# --- Database Setup ---
DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./.data/docuchat.db")
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False} if "sqlite" in DATABASE_URL else {})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    first_name = Column(String)
    last_name = Column(String)
    email = Column(String, unique=True, index=True)
    password_hash = Column(String)
    openrouter_key = Column(String, nullable=True)
    openrouter_model = Column(String, default="openai/gpt-4o")
    hf_key = Column(String, nullable=True)
    pinecone_key = Column(String, nullable=True)
    pinecone_index = Column(String, nullable=True)
    
    folders = relationship("Folder", back_populates="owner", cascade="all, delete-orphan")
    files = relationship("FileModel", back_populates="owner", cascade="all, delete-orphan")

class Folder(Base):
    __tablename__ = "folders"
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String)
    user_id = Column(Integer, ForeignKey("users.id"))
    
    owner = relationship("User", back_populates="folders")
    files = relationship("FileModel", back_populates="folder", cascade="all, delete-orphan")

class FileModel(Base):
    __tablename__ = "files"
    id = Column(Integer, primary_key=True, index=True)
    filename = Column(String)
    user_id = Column(Integer, ForeignKey("users.id"))
    folder_id = Column(Integer, ForeignKey("folders.id"), nullable=True)
    
    owner = relationship("User", back_populates="files")
    folder = relationship("Folder", back_populates="files")

Base.metadata.create_all(bind=engine)

# --- Auth Setup ---
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
SECRET_KEY = os.getenv("SECRET_KEY", "super-secret-key-for-local-dev-only")
ALGORITHM = "HS256"

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def get_current_user(request: Request, db: Session = Depends(get_db)):
    authorization = request.headers.get("Authorization")
    if not authorization or not authorization.startswith("Bearer "):
        print(f"[AUTH ERROR] Missing or invalid header: {authorization}")
        raise HTTPException(status_code=401, detail="Invalid token header")
    
    token = authorization.split(" ")[1]
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id: int = int(payload.get("sub"))
        if user_id is None:
            print("[AUTH ERROR] Missing 'sub' in token payload")
            raise HTTPException(status_code=401, detail="Invalid token payload")
    except ValueError:
        print("[AUTH ERROR] 'sub' claim is not a valid integer")
        raise HTTPException(status_code=401, detail="Invalid token subject")
    except Exception as e:
        print(f"[AUTH ERROR] PyJWT decode failed: {e}")
        raise HTTPException(status_code=401, detail="Token decoding failed")
        
    user = db.query(User).filter(User.id == user_id).first()
    if user is None:
        print(f"[AUTH ERROR] User ID {user_id} not found in DB")
        raise HTTPException(status_code=401, detail="User not found")
        
    return user

# --- Auth Endpoints ---
class AuthSignupRequest(BaseModel):
    first_name: str
    last_name: str
    email: str
    password: str

class AuthLoginRequest(BaseModel):
    email: str
    password: str

@app.post("/api/signup")
def signup(req: AuthSignupRequest, db: Session = Depends(get_db)):
    email = req.email.lower()
    if db.query(User).filter(User.email == email).first():
        raise HTTPException(status_code=400, detail="Email already exists")
    
    hashed_password = pwd_context.hash(req.password)
    user = User(first_name=req.first_name, last_name=req.last_name, email=email, password_hash=hashed_password)
    db.add(user)
    db.commit()
    db.refresh(user)
    
    token = jwt.encode({"sub": str(user.id), "exp": datetime.utcnow() + timedelta(days=7)}, SECRET_KEY, algorithm=ALGORITHM)
    if isinstance(token, bytes):
        token = token.decode("utf-8")
    return {"token": token, "email": user.email, "name": user.first_name}

@app.post("/api/login")
def login(req: AuthLoginRequest, db: Session = Depends(get_db)):
    email = req.email.lower()
    user = db.query(User).filter(User.email == email).first()
    if not user or not pwd_context.verify(req.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Incorrect email or password")
    
    token = jwt.encode({"sub": str(user.id), "exp": datetime.utcnow() + timedelta(days=7)}, SECRET_KEY, algorithm=ALGORITHM)
    if isinstance(token, bytes):
        token = token.decode("utf-8")
    return {"token": token, "email": user.email, "name": user.first_name}

# --- Settings Endpoints ---
class SettingsUpdate(BaseModel):
    openrouter_key: str
    openrouter_model: str
    hf_key: str
    pinecone_key: str
    pinecone_index: str

@app.get("/api/settings")
def get_settings(user: User = Depends(get_current_user)):
    return {
        "openrouter_key": user.openrouter_key or "",
        "openrouter_model": user.openrouter_model or "openai/gpt-4o",
        "hf_key": user.hf_key or "",
        "pinecone_key": user.pinecone_key or "",
        "pinecone_index": user.pinecone_index or ""
    }

@app.post("/api/settings")
def update_settings(req: SettingsUpdate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    user.openrouter_key = req.openrouter_key
    user.openrouter_model = req.openrouter_model
    user.hf_key = req.hf_key
    user.pinecone_key = req.pinecone_key
    user.pinecone_index = req.pinecone_index
    db.commit()
    return {"message": "Settings updated"}

# --- Folder & File Endpoints ---
class FolderCreate(BaseModel):
    name: str

@app.post("/api/folders")
def create_folder(req: FolderCreate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    folder = Folder(name=req.name, user_id=user.id)
    db.add(folder)
    db.commit()
    db.refresh(folder)
    return {"id": folder.id, "name": folder.name}

@app.get("/api/data")
def get_data(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    folders = db.query(Folder).filter(Folder.user_id == user.id).all()
    files = db.query(FileModel).filter(FileModel.user_id == user.id).all()
    
    return {
        "folders": [{"id": f.id, "name": f.name} for f in folders],
        "files": [{"id": f.id, "filename": f.filename, "folder_id": f.folder_id} for f in files]
    }

# --- File Processing ---
def extract_text_from_file(file: UploadFile, content: bytes) -> str:
    filename = file.filename.lower()
    text = ""
    try:
        if filename.endswith('.pdf'):
            pdf_reader = pypdf.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        elif filename.endswith('.docx'):
            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif filename.endswith('.csv'):
            df = pd.read_csv(io.BytesIO(content))
            text = df.to_string(index=False)
        elif filename.endswith('.txt'):
            text = content.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format.")
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error reading file: {str(e)}")
    return text

@app.post("/api/upload")
async def upload_document(
    file: UploadFile = File(...),
    folder_id: int = Form(None),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user)
):
    if not user.hf_key or not user.pinecone_key or not user.pinecone_index:
        raise HTTPException(status_code=400, detail="Missing API keys in settings.")
        
    try:
        content = await file.read()
        text = extract_text_from_file(file, content)
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted.")

        # Save File Record
        new_file = FileModel(filename=file.filename, user_id=user.id, folder_id=folder_id)
        db.add(new_file)
        db.commit()
        db.refresh(new_file)

        # Chunking
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_text(text)
        
        # Attach Metadata for Filtering
        documents = [
            Document(
                page_content=chunk, 
                metadata={
                    "source": file.filename,
                    "user_id": user.id,
                    "folder_id": folder_id or -1, # Use -1 for root files
                    "file_id": new_file.id
                }
            ) for chunk in chunks
        ]
        
        # Embeddings & Upsert
        embeddings = HuggingFaceEndpointEmbeddings(
            huggingfacehub_api_token=user.hf_key, 
            model="sentence-transformers/all-MiniLM-L6-v2"
        )
        os.environ["PINECONE_API_KEY"] = user.pinecone_key
        PineconeVectorStore.from_documents(documents, embeddings, index_name=user.pinecone_index)
        
        return {"message": "Success", "file_id": new_file.id}
        
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

class ChatRequest(BaseModel):
    message: str
    selected_folders: List[int] = []
    selected_files: List[int] = []

@app.post("/api/chat")
async def chat(request: ChatRequest, user: User = Depends(get_current_user)):
    if not user.openrouter_key or not user.pinecone_key or not user.pinecone_index:
        raise HTTPException(status_code=400, detail="Missing API keys in settings.")
        
    try:
        embeddings = HuggingFaceEndpointEmbeddings(
            huggingfacehub_api_token=user.hf_key, 
            model="sentence-transformers/all-MiniLM-L6-v2"
        )
        os.environ["PINECONE_API_KEY"] = user.pinecone_key
        vectorstore = PineconeVectorStore.from_existing_index(
            index_name=user.pinecone_index,
            embedding=embeddings
        )
        
        # Build Filter
        filter_dict = {"user_id": {"$eq": user.id}}
        
        if request.selected_files or request.selected_folders:
            conditions = []
            if request.selected_files:
                conditions.append({"file_id": {"$in": request.selected_files}})
            if request.selected_folders:
                conditions.append({"folder_id": {"$in": request.selected_folders}})
                
            # Combine logic: files OR folders
            if len(conditions) > 1:
                filter_dict["$or"] = conditions
            else:
                filter_dict.update(conditions[0])

        retriever = vectorstore.as_retriever(search_kwargs={"k": 5, "filter": filter_dict})
        docs = retriever.invoke(request.message)
        context = "\n\n".join([doc.page_content for doc in docs])
        
        llm = ChatOpenAI(
            model=user.openrouter_model or "openai/gpt-4o",
            openai_api_key=user.openrouter_key,
            openai_api_base="https://openrouter.ai/api/v1",
            max_tokens=1024,
            temperature=0.7
        )
        
        prompt = f"""You are a helpful AI assistant. Use the following context to answer the user's question. 
If the answer is not in the context, just answer based on your general knowledge, but mention that it's not from the uploaded documents.

Context:
{context}

Question:
{request.message}

Answer:"""
        
        response = llm.invoke(prompt)
        sources = list(set([doc.metadata.get("source", "Unknown") for doc in docs]))
        
        return {"answer": response.content, "sources": sources}
        
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

app.mount("/", StaticFiles(directory="public", html=True), name="public")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
